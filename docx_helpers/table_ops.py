"""
docx_helpers/table_ops.py
==========================
Helper untuk mengisi tabel docx secara dinamis: mencari tabel
berdasarkan teks header, lalu menggandakan baris template (master row)
sesuai jumlah data, tanpa merusak style (font/border) baris asli.
"""

import copy

from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def find_table_by_header(doc, header_keywords):
    for tbl in doc.tables:
        if len(tbl.rows) < 2:
            continue
        header_texts = [c.text.strip().lower() for c in tbl.rows[0].cells]
        if all(any(kw.lower() in h for h in header_texts) for kw in header_keywords):
            return tbl
    return None


def set_cell_text_preserve_style(cell, lines):
    """Mengisi teks pada sel sambil mempertahankan style run pertama
    yang sudah ada di template (font, ukuran, bold)."""
    para = cell.paragraphs[0]
    font_name, font_size, font_bold = "Arial", None, False
    if para.runs:
        r0 = para.runs[0]
        if r0.font.name:
            font_name = r0.font.name
        if r0.font.size:
            font_size = r0.font.size
        font_bold = bool(r0.font.bold)

    from docx.shared import Pt
    if font_size is None:
        font_size = Pt(11)

    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    for p in list(cell.paragraphs[1:]):
        p._element.getparent().remove(p._element)

    for idx, line in enumerate(lines):
        if idx > 0:
            run = para.add_run()
            run._element.append(OxmlElement('w:br'))
        run = para.add_run(line)
        run.font.name = font_name
        run.font.size = font_size
        run.bold = font_bold


def fill_table_rows_from_master(doc, header_keywords, rows_data):
    """Menggandakan baris ke-2 tabel (baris master/contoh) sebanyak
    jumlah rows_data, lalu mengisi tiap baris dengan data terkait."""
    tbl = find_table_by_header(doc, header_keywords)
    if tbl is None or not rows_data:
        return

    template_tr = tbl.rows[1]._tr
    parent = template_tr.getparent()
    new_trs = [copy.deepcopy(template_tr) for _ in rows_data]

    ref = template_tr
    for new_tr in new_trs:
        ref.addnext(new_tr)
        ref = new_tr
    parent.remove(template_tr)

    from docx.table import _Cell
    for new_tr, data_row in zip(new_trs, rows_data):
        cells_tc = new_tr.findall(qn('w:tc'))
        for tc, val in zip(cells_tc, data_row):
            cell = _Cell(tc, tbl)
            lines = val.split('\n') if isinstance(val, str) else [str(val)]
            set_cell_text_preserve_style(cell, lines)
