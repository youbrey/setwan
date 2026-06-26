"""
docx_helpers/combine.py
========================
Helper untuk menggabungkan beberapa file .docx menjadi satu dokumen
(setiap file menjadi satu halaman / beberapa halaman, dipisah page break).
Dipakai oleh hampir semua jenis surat yang punya banyak halaman per-orang
(SPD per orang, surat pemberitahuan per tujuan, daftar hadir per tujuan,
dan surat undangan paripurna per penerima).
"""

from docx import Document


def combine_word_pages(files_list, out_path):
    """Gabungkan beberapa file docx menjadi satu file, dipisah page break."""
    if not files_list:
        return
    try:
        master = Document(files_list[0])
        for f in files_list[1:]:
            master.add_page_break()
            subdoc = Document(f)
            for element in subdoc.element.body:
                if not element.tag.endswith('sectPr'):
                    master.element.body.append(element)
        if master.paragraphs:
            last_p = master.paragraphs[-1]
            if not last_p.text.strip():
                last_p._element.getparent().remove(last_p._element)
        master.save(out_path)
    except Exception as e:
        print(f"Gagal menggabungkan: {e}")
