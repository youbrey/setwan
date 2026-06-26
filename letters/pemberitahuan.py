"""
letters/pemberitahuan.py
==========================
Logika untuk SURAT PEMBERITAHUAN (bagian dari rangkaian surat
perjalanan dinas: surat tugas -> SPD -> surat pemberitahuan ke
instansi tujuan). Termasuk deteksi otomatis kategori & jumlah
pelaksana DPRD yang ditampilkan di surat pemberitahuan.

Modul ini terpisah dari letters/undangan.py (surat undangan rapat
paripurna), karena keduanya punya alur bisnis & data yang berbeda
sama sekali walaupun sama-sama "surat".
"""

import os
import re
import tempfile

from docx import Document
from docxtpl import DocxTemplate

from config import KATEGORI_DPRD_ORDER
from utils.nomor import increment_nomor
from utils.tanggal import generate_periods


def _label_kategori_dprd(cat, jabatan_list):
    if cat == "Pimpinan DPRD":
        return "Pimpinan DPRD"

    has_pimpinan = any(
        ("ketua" in j.lower() or "sekretaris" in j.lower()) for j in jabatan_list
    )
    has_anggota = any(
        ("anggota" in j.lower() and "ketua" not in j.lower()) for j in jabatan_list
    )

    if has_pimpinan and has_anggota:
        label = f"Pimpinan dan Anggota {cat}"
    elif has_pimpinan:
        label = f"Pimpinan {cat}"
    elif has_anggota:
        label = f"Anggota {cat}"
    else:
        label = cat
    if "DPRD" not in label:
        label += " DPRD"
    return label


def compute_pelaksana_dprd_summary(selected_dprd):
    """Mengelompokkan pelaksana DPRD per kategori, lalu menghasilkan
    daftar (label, jumlah) sesuai urutan KATEGORI_DPRD_ORDER."""
    by_cat = {}
    for p in selected_dprd:
        cat = str(p.get('kategori', '')).strip()
        by_cat.setdefault(cat, []).append(p.get('jabatan', ''))

    summary = []
    for cat in KATEGORI_DPRD_ORDER:
        jabatan_list = by_cat.get(cat)
        if not jabatan_list:
            continue
        label = _label_kategori_dprd(cat, jabatan_list)
        summary.append((label, len(jabatan_list)))

    for cat, jabatan_list in by_cat.items():
        if cat not in KATEGORI_DPRD_ORDER:
            label = _label_kategori_dprd(cat, jabatan_list)
            summary.append((label, len(jabatan_list)))
    return summary


def _remove_empty_pelaksana_lines(doc):
    for p in list(doc.paragraphs):
        txt = p.text
        normalized = re.sub(r'\s+', ' ', txt).strip()

        if "Pendamping ASN" in txt:
            m = re.search(r':\s*(\d+)\s+Orang', normalized)
            if m and int(m.group(1)) == 0:
                p._element.getparent().remove(p._element)
                continue

        if "Kota Bitung" in txt and "Orang" in txt:
            m = re.search(r':\s*(\d+)\s+Orang', normalized)
            if not m:
                p._element.getparent().remove(p._element)
                continue


def apply_pelaksana_dprd_summary_to_ctx(ctx, selected_dprd, max_slots=4):
    summary = compute_pelaksana_dprd_summary(selected_dprd)
    for i in range(1, max_slots + 1):
        if i <= len(summary):
            label, jumlah = summary[i - 1]
            ctx[f"pelaksana_tugas_{i}"] = label
            ctx[f"jlh_pelaksana_dprd{i}"] = jumlah
        else:
            ctx[f"pelaksana_tugas_{i}"] = ""
            ctx[f"jlh_pelaksana_dprd{i}"] = ""
    return ctx


def buat_surat_pemberitahuan_multi(template_path, ctx, selected_dprd, selected_asn,
                                    destinations, base_number, out_path,
                                    label_asn="Pendamping ASN"):
    """Membuat satu surat pemberitahuan multi-halaman, satu halaman per
    tujuan, dengan nomor surat yang otomatis bertambah per halaman."""
    periods = generate_periods(ctx.get("tanggal_mulai", ""), destinations)

    base_ctx = ctx.copy()
    apply_pelaksana_dprd_summary_to_ctx(base_ctx, selected_dprd)
    base_ctx["pelaksana_tugas_asn_info"] = label_asn
    base_ctx["jlh_pelaksana_asn"] = len(selected_asn)

    master_doc = None
    first = True

    for idx, period in enumerate(periods):
        nomor_surat = increment_nomor(base_number, idx)
        page_ctx = base_ctx.copy()
        page_ctx["nomor_surat_info"] = nomor_surat
        page_ctx["tujuan_surat_info"] = period["tujuan"]
        page_ctx["hari_info"] = period["hari"]
        page_ctx["tanggal_bertugas_info"] = period["tanggal"]

        tmp_docx = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
        doc_tpl = DocxTemplate(template_path)
        doc_tpl.render(page_ctx)
        doc_tpl.save(tmp_docx)

        doc = Document(tmp_docx)
        _remove_empty_pelaksana_lines(doc)
        doc.save(tmp_docx)

        if first:
            master_doc = Document(tmp_docx)
            first = False
        else:
            master_doc.add_page_break()
            subdoc = Document(tmp_docx)
            for element in subdoc.element.body:
                if not element.tag.endswith('sectPr'):
                    master_doc.element.body.append(element)

        os.unlink(tmp_docx)

    if master_doc:
        if master_doc.paragraphs:
            last_p = master_doc.paragraphs[-1]
            if not last_p.text.strip():
                last_p._element.getparent().remove(last_p._element)
        master_doc.save(out_path)
