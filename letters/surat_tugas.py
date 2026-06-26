"""
letters/surat_tugas.py
=======================
Logika khusus untuk SURAT TUGAS (perjalanan dinas), baik untuk DPRD
maupun ASN. Modul ini TIDAK menyentuh logika Surat Undangan -- lihat
letters/undangan.py untuk itu.

Surat Tugas punya 2 varian template:
- "biasa": jika jumlah pelaksana <= 3 orang (placeholder per-orang)
- "tabel": jika jumlah pelaksana > 3 orang (baris tabel digandakan)
"""

import re

from docx import Document
from docxtpl import DocxTemplate

from config import (
    TEMPLATE_ST_DPRD_BIASA, TEMPLATE_ST_DPRD_TABEL,
    TEMPLATE_ST_ASN_BIASA, TEMPLATE_ST_ASN_TABEL,
)
from docx_helpers.table_ops import fill_table_rows_from_master


# ---------------------------------------------------------------------------
# CLEANUP: hapus baris kosong "Nama:" / "Jabatan:" pada surat tugas biasa
# jika slot pelaksana tidak terisi penuh (misal hanya 2 dari 3 slot).
# ---------------------------------------------------------------------------
def cleanup_surat_tugas_biasa(doc, template_type):
    paragraphs = list(doc.paragraphs)
    to_remove = set()

    if template_type == 'dprd':
        i = 0
        while i < len(paragraphs):
            p = paragraphs[i]
            stripped = p.text.replace('\t', ' ').strip()
            if re.match(r'^Nama\s*:\s*$', stripped):
                to_remove.add(id(p))
                if i + 1 < len(paragraphs):
                    next_p = paragraphs[i + 1]
                    next_stripped = next_p.text.replace('\t', ' ').strip()
                    if re.match(r'^Jabatan\s*:\s*$', next_stripped):
                        to_remove.add(id(next_p))
                        i += 2
                        continue
            i += 1
    else:
        i = 0
        while i < len(paragraphs):
            p = paragraphs[i]
            stripped = p.text.replace('\t', ' ').strip()
            if re.match(r'^Nama\s*:\s*$', stripped):
                to_remove.add(id(p))
                j = i + 1
                while j < len(paragraphs):
                    ns = paragraphs[j].text.replace('\t', ' ').strip()
                    if re.match(r'^(Pangkat|NIP|Jabatan)\s*:\s*$', ns):
                        to_remove.add(id(paragraphs[j]))
                        j += 1
                    else:
                        break
                i = j
                continue
            i += 1

    for p in paragraphs:
        if id(p) in to_remove:
            p._element.getparent().remove(p._element)


# ---------------------------------------------------------------------------
# SURAT TUGAS DPRD
# ---------------------------------------------------------------------------
def buat_surat_tugas_dprd(ctx, selected_dprd, out_path):
    n = len(selected_dprd)
    render_ctx = ctx.copy()

    if n <= 3:
        template_path = TEMPLATE_ST_DPRD_BIASA
        for i in range(1, 4):
            p = selected_dprd[i - 1] if i <= n else {}
            render_ctx[f"pelaksana_tugas_{i}"] = p.get('nama', '')
            render_ctx[f"jabatan_pelaksana_{i}"] = p.get('jabatan', '')
        doc_tpl = DocxTemplate(template_path)
        doc_tpl.render(render_ctx)
        doc_tpl.save(out_path)
        doc = Document(out_path)
        cleanup_surat_tugas_biasa(doc, 'dprd')
        doc.save(out_path)
    else:
        template_path = TEMPLATE_ST_DPRD_TABEL
        render_ctx["loop"] = {"index": ""}
        render_ctx["tabel"] = {"nama": "", "jabatan": ""}
        doc_tpl = DocxTemplate(template_path)
        doc_tpl.render(render_ctx)
        doc_tpl.save(out_path)
        doc = Document(out_path)
        rows_data = [[str(i + 1), p.get('nama', ''), p.get('jabatan', '')]
                     for i, p in enumerate(selected_dprd)]
        fill_table_rows_from_master(doc, ["No", "Nama", "Jabatan"], rows_data)
        doc.save(out_path)


# ---------------------------------------------------------------------------
# SURAT TUGAS ASN
# ---------------------------------------------------------------------------
def buat_surat_tugas_asn(ctx, selected_asn, out_path):
    n = len(selected_asn)
    render_ctx = ctx.copy()

    if n <= 3:
        template_path = TEMPLATE_ST_ASN_BIASA
        for i in range(1, 4):
            p = selected_asn[i - 1] if i <= n else {}
            render_ctx[f"nama_asn_{i}"] = p.get('nama', '')
            render_ctx[f"pangkat_asn_{i}"] = p.get('pangkat', '')
            render_ctx[f"nip_asn_{i}"] = p.get('nip', '')
            render_ctx[f"jabatan_asn_{i}"] = p.get('jabatan', '')
        doc_tpl = DocxTemplate(template_path)
        doc_tpl.render(render_ctx)
        doc_tpl.save(out_path)
        doc = Document(out_path)
        cleanup_surat_tugas_biasa(doc, 'asn')
        doc.save(out_path)
    else:
        template_path = TEMPLATE_ST_ASN_TABEL
        render_ctx["loop"] = {"index": ""}
        render_ctx["tabel"] = {"nama_asn": "", "jabatan_asn": ""}
        doc_tpl = DocxTemplate(template_path)
        doc_tpl.render(render_ctx)
        doc_tpl.save(out_path)
        doc = Document(out_path)
        rows_data = []
        for i, p in enumerate(selected_asn):
            nama_col = f"{p.get('nama', '')}\nNIP. {p.get('nip', '-')}"
            jabatan_col = f"{p.get('jabatan', '-')}\n{p.get('pangkat', '-')}"
            rows_data.append([str(i + 1), nama_col, jabatan_col])
        fill_table_rows_from_master(doc, ["No", "Nama", "Jabatan"], rows_data)
        doc.save(out_path)
